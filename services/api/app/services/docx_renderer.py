"""
DOCX Renderer Service.
"""
from io import BytesIO
from typing import Any
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

logger = logging.getLogger(__name__)

class DocxRenderer:
    @staticmethod
    def render(resume_data: dict[str, Any]) -> bytes:
        """Render a TailoredResumeData JSON into a FAANG-style DOCX file."""
        # Minimal validation
        if not resume_data.get("fullName"):
            raise ValueError("Invalid Resume Data: fullName is required for DOCX rendering.")
            
        doc = docx.Document()
        
        # Set minimal margins for 1-pager (0.5 inches all around)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            
        # Font settings - standard FAANG choice is Arial/Calibri 10-11pt
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)
            
        # --- HEADER ---
        # Name
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_p.add_run(resume_data.get("fullName", "Candidate Name").upper())
        name_run.bold = True
        name_run.font.size = Pt(16)
        
        # Contact
        contact_line = resume_data.get("contactLine", "")
        if contact_line:
            contact_p = doc.add_paragraph()
            contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_p.add_run(contact_line)
            contact_run.font.size = Pt(10)

        # Helper to add section headers
        def add_section_header(title: str):
            doc.add_paragraph() # Spacing
            p = doc.add_paragraph()
            run = p.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(11)
            # Add bottom border equivalent by adding a line (hack for python-docx)
            # A real border requires oxml manipulation, so we just use underline for now
            # or a horizontal rule if possible. We will stick to bold caps for FAANG.
            p.paragraph_format.space_after = Pt(2)

        # --- SUMMARY ---
        summary = resume_data.get("summary", "")
        if summary:
            # For FAANG, summaries are rare, but if included they must be compact
            add_section_header("Professional Summary")
            sum_p = doc.add_paragraph(summary)
            sum_p.style.font.size = Pt(10)
            sum_p.paragraph_format.space_after = Pt(4)
            
        # --- EDUCATION ---
        education = resume_data.get("education", [])
        if education:
            add_section_header("Education")
            for edu in education:
                # Table for Left (School/Degree) and Right (Dates/Location)
                table = doc.add_table(rows=1, cols=2)
                table.autofit = True
                
                # School & Degree
                left_cell = table.cell(0, 0)
                p = left_cell.paragraphs[0]
                p.add_run(edu.get("school", "")).bold = True
                degree_text = f" — {edu.get('degree', '')}"
                if edu.get("gpa"):
                    degree_text += f" (GPA: {edu.get('gpa')})"
                p.add_run(degree_text)
                
                # Dates & Location
                right_cell = table.cell(0, 1)
                p2 = right_cell.paragraphs[0]
                p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                date_loc = f"{edu.get('location', '')} | {edu.get('dates', '')}"
                p2.add_run(date_loc.strip(" | "))
                
        # --- SKILLS ---
        skills = resume_data.get("skills", {})
        if skills:
            add_section_header("Technical Skills")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            
            def add_skill_line(label, value):
                if value:
                    p.add_run(f"{label}: ").bold = True
                    p.add_run(f"{value}\n")
                    
            add_skill_line("Languages", skills.get("languages", ""))
            add_skill_line("Frameworks", skills.get("frameworks", ""))
            add_skill_line("Cloud & DevOps", skills.get("cloudDevops", ""))
            add_skill_line("Databases", skills.get("databases", ""))

        # --- EXPERIENCE ---
        experience = resume_data.get("experience", [])
        if experience:
            add_section_header("Professional Experience")
            for exp in experience:
                table = doc.add_table(rows=1, cols=2)
                table.autofit = True
                
                left_cell = table.cell(0, 0)
                p = left_cell.paragraphs[0]
                p.add_run(exp.get("title", "")).bold = True
                p.add_run(f" | {exp.get('company', '')}")
                
                right_cell = table.cell(0, 1)
                p2 = right_cell.paragraphs[0]
                p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                date_loc = f"{exp.get('location', '')} | {exp.get('dates', '')}"
                p2.add_run(date_loc.strip(" | "))
                
                for bullet in exp.get("bullets", []):
                    b = doc.add_paragraph(bullet, style='List Bullet')
                    b.paragraph_format.left_indent = Inches(0.25)
                    b.paragraph_format.space_after = Pt(2)
                    b.style.font.size = Pt(10)

        # --- PROJECTS ---
        projects = resume_data.get("projects", [])
        if projects:
            add_section_header("Projects")
            for proj in projects:
                p = doc.add_paragraph()
                p.add_run(proj.get("name", "")).bold = True
                p.add_run(f" | {proj.get('tech', '')}")
                p.paragraph_format.space_after = Pt(2)
                
                for bullet in proj.get("bullets", []):
                    b = doc.add_paragraph(bullet, style='List Bullet')
                    b.paragraph_format.left_indent = Inches(0.25)
                    b.paragraph_format.space_after = Pt(2)
                    b.style.font.size = Pt(10)

        # Save to bytes
        io = BytesIO()
        doc.save(io)
        return io.getvalue()
